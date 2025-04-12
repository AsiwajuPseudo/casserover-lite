
"""
    Class Collector for retrieving data from documents from different file formats such as csv, excel, docx, pdf etc.
    Methods to note: 
        1. pdf_lines method collects text from a .pdf as individual lines assigned to their individual pages.
        2. docx_paragraph method for collecting text from .docx files as paragaraphs
        3. docx_lines method for collecting text from .docx files as lines
        4. docx_style method for collecting text from .docx files with the text's respective style.
        5. pdf_raw method for collecting text from .pdf file in pages of the PDF.

"""

import csv
import json
import io
import os
import fitz
import docx
import time
import random
import openpyxl
import mammoth
from bs4 import BeautifulSoup
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from langchain.text_splitter import TokenTextSplitter

class Collector:
    @staticmethod
    def collect_csv(filename):
        try:
            data = []
            with open(filename, 'r', newline='') as csvfile:
                reader = csv.DictReader(csvfile)
                for row in reader:
                    data.append({key: row[key] for key in row})
            return data
        except FileNotFoundError:
            print(f"Error: File '{filename}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def find_arrays(d):
        array=[]
        for key, value in d.items():
            if isinstance(value, list):
                array=value
                break
            elif isinstance(value, dict):
                array = find_arrays(value)
        return array

    @staticmethod
    def collect_json(filename):
        try:
            with open(filename, 'r') as jsonfile:
                data = json.load(jsonfile)
                updated_data=Collector.find_arrays(data)
            return updated_data
        except FileNotFoundError:
            print(f"Error: File '{filename}' not found.")
            return []
        except json.JSONDecodeError:
            print(f"Error: File '{filename}' is not a valid JSON.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def collect_xlsx(filename):
        try:
            data = []
            workbook = openpyxl.load_workbook(filename)
            sheet = workbook.active
            
            headers = [cell.value for cell in sheet[1]]
            for row in sheet.iter_rows(min_row=2, values_only=True):
                data.append({headers[i]: row[i] for i in range(len(headers))})
            
            return data
        except FileNotFoundError:
            print(f"Error: File '{filename}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def pdf_lines(path):
        try:
            #open the file
            pdf_document = fitz.open(path)
            document = []
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text()
                lines = text.split('\n')
                page_lines = [{'n': i + 1, 'text': line} for i, line in enumerate(lines)]
                document.append({'page': page_num + 1, 'lines': page_lines})

            return document
        except FileNotFoundError:
            print(f"Error: File '{path}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def docx_paragraph(path):
        try:
            #open the file
            docx_file = docx.Document(path)
            paragraphs = [{'text':paragraph.text} for paragraph in docx_file.paragraphs]

            return paragraphs
        except FileNotFoundError:
            print(f"Error: File '{path}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def docx_lines(path):
        try:
            #open the file
            docx_file = docx.Document(path)
            paragraphs = [paragraph.text for paragraph in docx_file.paragraphs]
            n=1
            i=0
            num=1
            doc=[]
            page=[]
            for p in paragraphs:
                if n<12 and i<len(paragraphs)-1:
                    page.append({'n':n,'text':p})
                    n=n+1
                elif i==len(paragraphs)-1:
                    page.append({'n':n,'text':p})
                    doc.append({'page':num,'lines':page})
                else:
                    doc.append({'page':num,'lines':page})
                    page=[]
                    page.append({'n':n,'text':p})
                    n=1
                    num=num+1
                i=i+1

            return doc
        except FileNotFoundError:
            print(f"Error: File '{path}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def docx_styles(path):
        # Open the document
        doc = docx.Document(path)
        
        # List to store paragraphs with their styles and text
        paragraphs_with_styles = []
        
        # Iterate through each paragraph in the document
        for para in doc.paragraphs:
            ident=0
            if para.paragraph_format.left_indent is not None:
                ident=para.paragraph_format.left_indent.inches
            style=para.style.name
            if para.style.name=='Heading 1':
                style='h1'
            elif para.style.name=='Heading 2':
                style='h2'
            elif para.style.name=='Heading 3':
                style='h3'
            elif para.style.name=='Heading 4':
                style='h4'
            else:
                pass
            paragraph_data = {
                'style': style,
                'text': para.text,
                'ident':ident
            }
            paragraphs_with_styles.append(paragraph_data)
        
        return paragraphs_with_styles

    @staticmethod
    def docx_to_html(path):
        with open(path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
        
        return html

    @staticmethod
    def html_styles(path):
        try:
            with open(path, 'r', encoding='utf-8') as file:
                html_content = file.read()

            soup = BeautifulSoup(html_content, 'html.parser')
            parts = soup.find_all('div')
            if not parts:
                return []  # Return empty list if 'Section0' div is not found

            elements_with_styles = []
            for part in parts:
                for element in part.find_all(recursive=False): #only immediate children
                    if element.name:
                        style = element.name  # Use tag name as style equivalent
                        text = element.get_text(strip=False)
                        ident = 0.0

                        # Extract left margin indent from style attribute
                        if element.has_attr('style'):
                            style_str = element['style']
                            if 'margin-left:' in style_str:
                                margin_left = style_str.split('margin-left:')[1].split(';')[0].strip()
                                if 'px' in margin_left:
                                    try:
                                        ident = float(margin_left.replace('px', '')) / 96.0 # approximate conversion from px to inches (96 dpi)
                                    except ValueError:
                                        ident = 0.0 # handle non-numeric margin values
                                elif 'in' in margin_left:
                                  try:
                                    ident = float(margin_left.replace('in', ''))
                                  except ValueError:
                                    ident = 0.0
                                elif 'pt' in margin_left:
                                  try:
                                    ident = float(margin_left.replace('pt', '')) / 72.0 # approximate conversion from pt to inches (72 dpi)
                                  except ValueError:
                                    ident = 0.0

                        element_data = {
                            'style': style,
                            'text': text.replace('\xa0',' '),
                            'ident': ident
                        }
                        elements_with_styles.append(element_data)

            return elements_with_styles

        except FileNotFoundError:
            print(f"Error: File not found at {path}")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def pdf_raw(path):
        try:
            #open the file
            pdf_document = fitz.open(path)
            dataset = []
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text= page.get_text()
                dataset.append({'page_number':page_num,'text':text})

            return dataset
        except FileNotFoundError:
            print(f"Error: File '{path}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

# Example usage:
# collector = Collector()
# collected_data = collector.collect_csv('data.csv')
# print(collected_data)