import os
import json
from gpt import GPT
from collector import Collector
from langchain.text_splitter import TokenTextSplitter
import os
import random
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX

class Heads:
    def __init__(self, euclid):
        self.euclid = euclid
        self.gpt = GPT()
        self.items = []

        self.drafter = """
        You are part of an AI agent being used for drafting heads of arguments after legal research for a lawyer in Zimbabwe.
        The agent receives a prompt from a user and it uses an Large Language Model(LLM) to generate search phrases that are 
        used to retrieve documents from a vector database using cosine search. Each and everyone of the retrieved document is then
        sent to an LLM that extract and generate an analysis/review/answer of the document's content in relation to the search phrases
        and the users prompt. You are the last step in the Agent's process, you are provided with the analysis/reviews/answers from the
        previous step and your role is to generate heads of arguments that use the analysis and the lawyer's data in the chat history.
        Generate actual and full heads of arguments using the data you are provided with, do not generate anything other than that.
        Your response should be long as much as possible and you should include all the details (in the best way possible). You should also provide citations for everything and
        cite case law and legislations (this is mandatory) without hallucinationsof cource.
        The structure of your JSON answer should be as follows: {'answer':[....]}.
        The array in the answer is a list of sections in your answer. A section can be a header,paragraph, list, table, map, pie chart or bar chat
        Every section type follows a certain format described below and your response should be of that format only.
        1. header- used to represent the start of certain content. Format={'type':'header','data':header text}
        2. paragraph- used to show a paragraph of text. Format= {'type':'paragraph','data':paragraph text}
        3. list- used to represent a list of items. Format= {'type':'list','data':[...items]} where items is an array of text.
            Normally, a list should be preceded by a header section which will be the name of the list
        4. table- used to represent a table of data. Format={'type':'table','data':{'columns':[...columns],'values':[...values]},}
            columns is an array of the columns contained in the table and every column should be of the format:
                {'title':name of column,'dataIndex':name of index to target the value in the row values,'key':key to target the value in the row values}
            values are an array of rows in the table and every row should be of key-value pairs matching the columns contained like {'':...,'':...}.


        SOME NOTES
        a. Your answer should be proper json data.
        b. Do not put 'None' or 'Null' or blank where a numerical value is required, if there is no sufficient numeriacal data provided then put your own estimates appropriately.
        c. If the provided research is not relevant to what the user wants, use your knowledge and if you have knowledge about the issue then generate the answer with such knowledge.
        d. Do not put an new lines (\n) or tabs (\t) or anything of that nature in your response.
        """

        self.researcher = """
        You are part of an AI Agent for legal research in Zimbabwe. Your role is to generate a research answer to a research topic(s)
        or question(s). You are provided with a document or part of a document and you have to go through the document to pick up all the 
        content in the document that is relevant to the research topic or topics and generate an accurate answer, review or analysis of that
        content in relation to the topic(s). Your response will be used by another LLM part of the agent to combine with other answers to
        other research questions to generate one final answer hence if there is relevant content you should make sure to mention it in
        your analysis, review or answer comprehensively.If there is no relevant content in the document then your response should be 'None',
        do not generate any response in such case that is long. Provide every necessary citation and cite the document itself if necessary
        and citations are very mandatory. Cite as much as possible without hallucinations and also cite precedent/case law used if the
        content is a court ruling and cite sections/subsections/paragraphs ect. if the provided document is alegislation or contract.
        """

        self.data_gather = """
        You are part of an AI-Agent for drafting heads of arguments for lawyers in an AI-powered research tool for lawyers in 
        Zimbabwe. Your responsibility is to generate search phrases that should to query a table from the vector database which
        contains all caselaw, statutes and regulations documents. However, before you can start generating any search phrase, you
        should be able to establish that the details which you need are provided by the user. This means if there is not enough 
        data to generate comprehensive search phrases (research questions) then you should instead ask the user to provide any
        further details which you will need. Your response should be in json format of the structure {'result':the result,'data':
        the data}. The result should be 'incomplete' if the provided details/ data from your interaction with the user does not
        have all the details you might need to generate the search phrases and in that case, the data key should have a question
        which should be asked to the user requesting them to provide any further details needed. When the details are all available
        the result should be 'complete' and the data key should be returned with an array containing a list of all the search
        phrases that are needed for searching and generating the heads of arguments. The search phrase item in the array should
        be a dict of strucuture {'phrase':the search phrase,'table': name of table to search}. The search phrase should be a
        statement able to generate accurate results in vector database cosine similarity search and it must be something relevant
        to searching for precedent or searching for sections of a statute or regulation. The names of the vector database tables
        are:
        """

    def gather(self, prompt, history, tables):
        messages = [{"role": "system", "content": self.data_gather + str(tables)}]
        for message in history:
            messages.append({"role": "user", "content": message['user']})
            messages.append({"role": "assistant", "content": str(message['system'])})
        messages.append({"role": "user", "content": prompt})
        answ = self.gpt.json_gpt(messages, 4060)
        answer = json.loads(answ)
        return answer

    def load_unique(self, data):
        unique_docs = set()
        for item in data:
            unique_docs.add((item['citation'], item['table'], item['table_id'], item['file_id'], item['filename']))
        sources = [
            {'citation': citation, 'table': table, 'table_id': table_id, 'file_id': file_id, 'filename': filename}
            for citation, table, table_id, file_id, filename in unique_docs
        ]
        return sources

    def open_file(self, file_id, filename, table, table_id):
        file_path = f'../temp/{table}-{table_id}/{file_id}-{filename}'
        collect = Collector()
        if filename.lower().endswith('.pdf'):
            document = collect.pdf_raw(file_path)
        elif filename.lower().endswith('.docx'):
            document = collect.docx_styles(file_path)
        elif filename.lower().endswith(('.htm', '.html')):
            document = collect.html_styles(file_path)
        return document

    def research(self, question, source):
        document = self.open_file(source['file_id'], source['filename'], source['table'], source['table_id'])
        text = '\n'.join(t['text'] for t in document)
        splitter = TokenTextSplitter(chunk_size=114000, chunk_overlap=200)
        chunks = splitter.split_text(text)
        answers = ''
        for chunk in chunks:
            context = f"Research topic (Question): {question}\n\n\n Document Name(Citation): {source['citation']}\nDocument/ Part of Document: {chunk}"
            messages = [{"role": "system", "content": self.researcher}]
            messages.append({"role": "user", "content": context})
            answers += '\n' + self.gpt.gpt_4o(messages, 12000)
        return answers

    def drafting(self, prompt, history, phrases, k):
        raw_sources = []
        for phrase in phrases:
            raw_sources.extend(self.euclid.search(phrase['table'], phrase['phrase'], k))

        sources = self.load_unique(raw_sources)
        researches = [
            {'citation': source['citation'], 'research_answer': self.research(str(phrases), source)}
            for source in sources
        ]

        context = f"Research: {researches}\n Prompt: {prompt}"
        messages = [{"role": "system", "content": self.drafter}]
        for message in history:
            messages.append({"role": "user", "content": message['user']})
            messages.append({"role": "assistant", "content": str(message['system'])})
        messages.append({"role": "user", "content": context})
        answ = self.gpt.json_gpt(messages, 16380)
        answer = json.loads(answ)
        answer['research'] = researches
        answer['phrases'] = phrases
        answer['citations'] = sources
        return answer, sources

    def create_docx(self, data):
        try:
            if not os.path.exists('../documents_created/'):
                os.makedirs('../documents_created/')
            
            document = Document()
            
            for item in data:
                element_type = item.get("type")
                element_data = item.get("data")

                if element_type == "header":
                    document.add_heading(element_data, level=1)
                elif element_type == "paragraph":
                    document.add_paragraph(element_data)
                elif element_type == "list":
                    if isinstance(element_data, list):
                        for list_item in element_data:
                            document.add_paragraph(list_item, style="List Bullet")
                    else:
                        print(f"Warning: List data is not a list: {element_data}")
                else:
                    print(f"Warning: Unknown element type: {element_type}")
            
            filename = str(random.randint(100000, 9999999)) + '-doc.docx'
            filepath = os.path.join('../documents_created/', filename)
            document.save(filepath)
            
            return filename
        except Exception as e:
        	print(str(e))
        	return ''

    def run(self,tables, prompt, history, k=3):
        check = self.gather(prompt, history, tables)
        if check['result'] == 'incomplete':
        	answer={'answer':[{'type': 'paragraph', 'data': check['data']}]}
        	ans=json.dumps(answer)
        	return ans,[]
        else:
        	answer,sources=self.drafting(prompt, history, check['data'], k)
        	doc1=self.create_docx(answer['answer'])
        	answer1={'answer':[{'type':'paragraph','data':'Successfully created heads of arguments'},{'type':'document','data':doc1}]}
        	answer2=json.dumps(answer1)
        	return answer2,sources